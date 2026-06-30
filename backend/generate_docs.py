"""Generate all .docx documentation files for AutoRemediate AI.

Run:
  python generate_docs.py                     # all phase-1 docs (no scan data needed)
  python generate_docs.py --scan-id 5         # include real scan findings
  python generate_docs.py --verify-diff docs/verification_diff.json

Requires: python-docx
"""

import argparse
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── paths ─────────────────────────────────────────────────────────────────────

DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")

TEAM = [
    ("Marc Wael", "Team Lead"),
    ("Mohey Elden Mostafa", "Developer"),
    ("Jerome Arsany", "Developer"),
    ("Noureen Elsayed", "Developer"),
    ("Hassan Ibrahim", "Developer"),
    ("Philopater Helmy", "Developer"),
]
EVALUATOR = "Eng. Ahmed Attia"
PROJECT_TITLE = "AutoRemediate AI — Autonomous Vulnerability Remediation & Incident Response Hub"
TODAY = datetime.now().strftime("%B %d, %Y")


# ── helpers ───────────────────────────────────────────────────────────────────

def _header(doc: Document, title: str, subtitle: str = "") -> None:
    """Add standard project header."""
    h = doc.add_heading(PROJECT_TITLE, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading(title, level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {TODAY}  |  Evaluator: {EVALUATOR}")
    doc.add_paragraph(
        "Team: " + " · ".join(f"{n} ({r})" for n, r in TEAM)
    )
    doc.add_paragraph("")


def _save(doc: Document, filename: str) -> str:
    os.makedirs(DOCS_DIR, exist_ok=True)
    path = os.path.join(DOCS_DIR, filename)
    doc.save(path)
    print(f"  Saved → {path}")
    return path


def _table_row(table, *cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row


# ── 1. Tool Configuration Documentation ───────────────────────────────────────

def gen_tool_config():
    doc = Document()
    _header(doc, "Tool Configuration Documentation")

    doc.add_heading("1. Purpose", level=3)
    doc.add_paragraph(
        "This document describes the tools, APIs, and credential management strategy "
        "used in the AutoRemediate AI project."
    )

    doc.add_heading("2. Tools Used", level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light List Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Tool"; hdr[1].text = "Version/URL"; hdr[2].text = "Purpose"
    _table_row(tbl, "Nessus Professional", "https://192.168.244.129:8834", "Vulnerability scanner")
    _table_row(tbl, "Metasploitable2", "192.168.244.128", "Target VM")
    _table_row(tbl, "Supabase (PostgreSQL)", "db.ezmzethdxnhyxfnrxoky.supabase.co", "Audit database")
    _table_row(tbl, "Python 3.11 (Anaconda)", "conda env: autoremediate", "Engine runtime")
    _table_row(tbl, "Paramiko", "≥3.4", "SSH remediation transport")
    _table_row(tbl, "Requests", "≥2.31", "Nessus API client")
    doc.add_paragraph("")

    doc.add_heading("3. Credential & Secret Management", level=3)
    doc.add_paragraph(
        "SECURITY DECISION: Real credentials are stored only in backend/config.json, which "
        "is listed in .gitignore and is NEVER committed to the repository. "
        "The committed file backend/config.example.json contains placeholder text so that "
        "contributors know the expected structure without exposing secrets."
    )
    doc.add_paragraph(
        "config.json holds Nessus API keys, SSH credentials for the target, and the "
        "Supabase connection string. Each developer creates their own config.json from "
        "config.example.json."
    )

    doc.add_heading("4. Nessus API Authentication", level=3)
    doc.add_paragraph(
        "All Nessus API calls use the X-ApiKeys HTTP header:\n"
        "  X-ApiKeys: accessKey=<ACCESS_KEY>; secretKey=<SECRET_KEY>\n\n"
        "SSL certificate verification is disabled (verify=False) because Nessus uses "
        "a self-signed certificate on port 8834."
    )

    doc.add_heading("5. Target Upstart Compatibility", level=3)
    doc.add_paragraph(
        "Metasploitable2 runs Ubuntu 8.04 (Hardy Heron) with Upstart, not systemd. "
        "All service management commands use:\n"
        "  sudo service <name> stop / start / status\n\n"
        "systemctl is NOT used and will fail silently on this target. "
        "This is enforced throughout remediator.py."
    )

    return _save(doc, "Tool_Configuration_Documentation.docx")


# ── 2. Assessment Scope Document ──────────────────────────────────────────────

def gen_assessment_scope():
    doc = Document()
    _header(doc, "Assessment Scope Document")

    doc.add_heading("1. Scope Overview", level=3)
    doc.add_paragraph(
        "This penetration and vulnerability assessment targets a controlled lab environment "
        "consisting of a single intentionally-vulnerable virtual machine (Metasploitable2). "
        "All activity is authorised within the DEPI Capstone lab network."
    )

    doc.add_heading("2. In-Scope Assets", level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light List Accent 1"
    h = tbl.rows[0].cells
    h[0].text = "Asset"; h[1].text = "IP"; h[2].text = "OS"; h[3].text = "Role"
    _table_row(tbl, "Metasploitable2", "192.168.244.128", "Ubuntu 8.04 (Upstart)", "Target — intentionally vulnerable VM")
    _table_row(tbl, "Kali Linux (Engine)", "192.168.244.129", "Kali Rolling", "Scanner / Remediator host")
    doc.add_paragraph("")

    doc.add_heading("3. Out-of-Scope", level=3)
    for item in [
        "Any host on the external internet",
        "Production systems or real infrastructure",
        "Denial-of-service testing",
        "Physical security testing",
    ]:
        doc.add_paragraph(f"• {item}")

    doc.add_heading("4. Scan Parameters", level=3)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Light List Accent 1"
    tbl2.rows[0].cells[0].text = "Parameter"
    tbl2.rows[0].cells[1].text = "Value"
    for k, v in [
        ("Scan Template", "Basic Network Scan"),
        ("Target Ports", "All (1-65535 default)"),
        ("Scan User", "Unauthenticated (external perspective)"),
        ("Scan Window", "Week 1–2 of capstone project"),
        ("Approval Required", "Yes — human-in-the-loop for all Critical/High"),
    ]:
        _table_row(tbl2, k, v)

    doc.add_heading("5. Rules of Engagement", level=3)
    for rule in [
        "No remediation shall execute without explicit operator keystroke approval.",
        "SSH (port 22) must never be stopped or network-blocked.",
        "All remediation activity is logged to Supabase with timestamp and operator decision.",
        "Retry limit: 3 attempts per finding before marking Failed.",
        "Post-remediation verification requires TWO independent checks (on-target + external socket).",
    ]:
        doc.add_paragraph(f"• {rule}")

    return _save(doc, "Assessment_Scope_Document.docx")


# ── 3. Vulnerability Scan Report ──────────────────────────────────────────────

def gen_vuln_scan_report(findings=None, scan_id=None):
    doc = Document()
    _header(doc, "Vulnerability Scan Report",
            f"Nessus Basic Network Scan — Metasploitable2 (192.168.244.128)"
            + (f" — Scan ID {scan_id}" if scan_id else ""))

    doc.add_heading("1. Executive Summary", level=3)
    if findings:
        critical = sum(1 for f in findings if f["severity"] == "Critical")
        high = sum(1 for f in findings if f["severity"] == "High")
        medium = sum(1 for f in findings if f["severity"] == "Medium")
        low = sum(1 for f in findings if f["severity"] == "Low")
        doc.add_paragraph(
            f"The Nessus scan identified {len(findings)} vulnerabilities: "
            f"{critical} Critical, {high} High, {medium} Medium, {low} Low."
        )
    else:
        doc.add_paragraph(
            "[PENDING — to be populated with real Nessus scan results after scan completion]"
        )

    doc.add_heading("2. Findings Table", level=3)
    if findings:
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light List Accent 1"
        h = tbl.rows[0].cells
        for i, label in enumerate(["#", "Plugin Name", "Severity", "Port", "CVE", "Score"]):
            h[i].text = label
        for idx, f in enumerate(findings, 1):
            _table_row(tbl, idx, f["plugin_name"], f["severity"],
                       f"{f['port']}/{f['protocol']}", f.get("cve_reference", "N/A"),
                       f.get("priority_score", ""))
    else:
        doc.add_paragraph("[Findings table will be populated after scan completes]")

    doc.add_heading("3. Notable Vulnerabilities", level=3)
    notable = [
        ("vsftpd 2.3.4 Backdoor (CVE-2011-2523)", "Port 21/tcp",
         "Remote code execution via hardcoded backdoor trigger sequence."),
        ("UnrealIRCd Backdoor (CVE-2010-2075)", "Port 6667/tcp",
         "IRC daemon contains intentional backdoor allowing arbitrary command execution."),
        ("Samba MS-RPC Shell Command Injection (CVE-2007-2447)", "Port 139/445",
         "Samba 3.0.20-3.0.25rc3 allows unauthenticated RCE via MS-RPC calls."),
        ("distccd RCE (CVE-2004-2687)", "Port 3632/tcp",
         "distccd allows arbitrary command execution without authentication."),
        ("Bind Shell Backdoor Detection", "Port 1524/tcp",
         "A bind shell listening on port 1524 provides unauthenticated root access."),
    ]
    for name, port_info, desc in notable:
        doc.add_paragraph(f"• {name} ({port_info}): {desc}")

    return _save(doc, "Vulnerability_Scan_Report.docx")


# ── 4. Initial Analysis Document ──────────────────────────────────────────────

def gen_initial_analysis(findings=None):
    doc = Document()
    _header(doc, "Initial Analysis Document")

    doc.add_heading("1. Attack Surface Summary", level=3)
    doc.add_paragraph(
        "Metasploitable2 exposes an unusually large attack surface due to its intentional "
        "misconfiguration. Multiple services have known backdoors, no authentication requirements, "
        "and publicly available exploit modules in Metasploit Framework."
    )

    doc.add_heading("2. Most Severe Findings", level=3)
    if findings:
        top = [f for f in findings if f["severity"] in ("Critical", "High")][:10]
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light List Accent 1"
        h = tbl.rows[0].cells
        for i, label in enumerate(["Plugin", "Severity", "Port", "CVE", "Exploit Available"]):
            h[i].text = label
        for f in top:
            _table_row(tbl, f["plugin_name"], f["severity"],
                       f"{f['port']}/{f['protocol']}", f.get("cve_reference", "N/A"),
                       "Yes" if f.get("exploit_available") else "No")
    else:
        doc.add_paragraph("[Populated after scan completion]")

    doc.add_heading("3. Service Risk Matrix", level=3)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Light List Accent 1"
    h = tbl2.rows[0].cells
    for i, label in enumerate(["Service", "Port", "Risk Level", "Notes"]):
        h[i].text = label
    for svc, port, risk, note in [
        ("vsftpd 2.3.4", "21/tcp", "Critical", "Known backdoor trigger: user :)"),
        ("UnrealIRCd 3.2.8.1", "6667/tcp", "Critical", "Hardcoded backdoor in AB prefix"),
        ("Samba 3.0.20", "139/445/tcp", "Critical", "usermap_script RCE"),
        ("distccd", "3632/tcp", "High", "RCE without auth"),
        ("Bind shell", "1524/tcp", "Critical", "Direct root shell no auth"),
        ("Telnet", "23/tcp", "High", "Plaintext credential transmission"),
        ("SSH", "22/tcp", "Low", "Known version info; keep running"),
        ("VNC", "5900/tcp", "High", "No auth or weak password"),
        ("Tomcat", "8180/tcp", "Medium", "Default credentials possible"),
        ("NFS", "2049/tcp", "Medium", "World-writable exports"),
    ]:
        _table_row(tbl2, svc, port, risk, note)

    doc.add_heading("4. Recommended Remediation Priority", level=3)
    doc.add_paragraph(
        "Priority order (by AutoRemediate scoring algorithm):\n"
        "1. Bind shell backdoor (port 1524) — Critical + no auth + confirmed open\n"
        "2. vsftpd backdoor (port 21) — Critical + Metasploit module\n"
        "3. UnrealIRCd backdoor (port 6667) — Critical + Metasploit module\n"
        "4. Samba usermap_script (port 445) — Critical + Metasploit module\n"
        "5. distccd RCE (port 3632) — High + Metasploit module"
    )

    return _save(doc, "Initial_Analysis_Document.docx")


# ── 5. Prioritization Report ──────────────────────────────────────────────────

def gen_prioritization_report(findings=None):
    doc = Document()
    _header(doc, "Prioritization Report")

    doc.add_heading("1. Scoring Methodology", level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light List Accent 1"
    tbl.rows[0].cells[0].text = "Criterion"
    tbl.rows[0].cells[1].text = "Points"
    for criterion, pts in [
        ("Critical severity", "+100"),
        ("High severity", "+70"),
        ("Medium severity", "+40"),
        ("Low severity", "+10"),
        ("Known backdoor / Metasploit module / actively exploited", "+40"),
        ("Public PoC (not weaponised)", "+20"),
        ("No authentication required", "+15"),
        ("Service confirmed listening (port > 0)", "+15"),
    ]:
        _table_row(tbl, criterion, pts)
    doc.add_paragraph("")

    doc.add_heading("2. Ranked Findings", level=3)
    if findings:
        from recommender import prioritise_and_recommend
        ranked = prioritise_and_recommend([f.copy() for f in findings])
        tbl2 = doc.add_table(rows=1, cols=6)
        tbl2.style = "Light List Accent 1"
        h = tbl2.rows[0].cells
        for i, label in enumerate(["Rank", "Plugin", "Severity", "Port", "Score", "Recommended Action"]):
            h[i].text = label
        for rank, f in enumerate(ranked, 1):
            preferred = f.get("recommendation", {}).get("preferred", {})
            _table_row(tbl2, rank, f["plugin_name"], f["severity"],
                       f"{f['port']}/{f['protocol']}", f["priority_score"],
                       preferred.get("description", ""))
    else:
        doc.add_paragraph("[Populated after scan completion and scoring]")

    return _save(doc, "Prioritization_Report.docx")


# ── 6. Remediation Plan ───────────────────────────────────────────────────────

def gen_remediation_plan(findings=None):
    doc = Document()
    _header(doc, "Remediation Plan")

    doc.add_heading("1. Approval Workflow", level=3)
    for step in [
        "Engine scores all findings and sorts Critical/High by priority score.",
        "For each finding, a Decision Card is displayed showing: recommended action, rationale, alternatives, and monitor-only option.",
        "Operator is REQUIRED to type one of: [A] Approve, [2-4] Alternative, [R] Reject, [Q] Quit.",
        "NO action executes without an explicit keypress — there is no auto-fire path.",
        "Every decision (Approve / Alternative / Reject) is written to Remediation_Activity_Logs with timestamp.",
    ]:
        doc.add_paragraph(f"  {step}")

    doc.add_heading("2. Preference Ladder", level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light List Accent 1"
    tbl.rows[0].cells[0].text = "Rung"
    tbl.rows[0].cells[1].text = "Action"
    tbl.rows[0].cells[2].text = "When Used"
    for rung, action, when in [
        ("1", "Patch/upgrade", "Root-cause fix — generally infeasible on Metasploitable2"),
        ("2", "Stop/disable service", "Backdoored or abandoned services (vsftpd, UnrealIRCd, etc.)"),
        ("3", "Harden/reconfigure", "Service must remain up; restrict access or disable dangerous options"),
        ("4", "iptables DROP", "Network-layer containment — standing alternative for any port"),
        ("5", "Monitor only", "Log only, no change to target — safe fallback always available"),
    ]:
        _table_row(tbl, rung, action, when)
    doc.add_paragraph("")

    doc.add_heading("3. SSH Safety Rule", level=3)
    doc.add_paragraph(
        "SSH (port 22 / service ssh/sshd) is NEVER stopped or firewall-blocked. "
        "Enforcement is in code (assert_not_ssh_kill in recommender.py and remediator.py). "
        "For SSH findings, the preferred action is always Rung 3 Harden: "
        "disable root login and restart the daemon."
    )

    doc.add_heading("4. Per-Finding Remediation Table", level=3)
    if findings:
        from recommender import prioritise_and_recommend
        ranked = prioritise_and_recommend([f.copy() for f in findings
                                          if f["severity"] in ("Critical", "High")])
        tbl2 = doc.add_table(rows=1, cols=5)
        tbl2.style = "Light List Accent 1"
        h = tbl2.rows[0].cells
        for i, label in enumerate(["Finding", "Preferred (Rung)", "Command", "Alternative", "Monitor Option"]):
            h[i].text = label
        for f in ranked:
            rec = f.get("recommendation", {})
            preferred = rec.get("preferred", {})
            alts = rec.get("alternatives", [])
            alt_text = alts[0]["description"] if alts else "iptables DROP"
            _table_row(tbl2,
                       f["plugin_name"],
                       f"Rung {preferred.get('rung','')} — {preferred.get('action','')}",
                       preferred.get("command", ""),
                       alt_text,
                       "Log only, no change")
    else:
        doc.add_paragraph("[Populated after scan completion]")

    return _save(doc, "Remediation_Plan.docx")


# ── 7. Verification Report ────────────────────────────────────────────────────

def gen_verification_report(diff=None, logs=None):
    doc = Document()
    _header(doc, "Verification Report")

    doc.add_heading("1. Methodology", level=3)
    doc.add_paragraph(
        "Post-remediation verification uses two independent checks:\n"
        "(i)  On-target: 'service <name> status' and 'netstat -tlnp | grep :<port>' "
        "confirm the service/port is stopped.\n"
        "(ii) External: a raw TCP socket probe from the engine host verifies the port "
        "is now connection-refused."
    )

    doc.add_heading("2. Before/After Comparison", level=3)
    if diff:
        doc.add_paragraph(f"Baseline Scan ID: {diff.get('baseline_scan_id')}")
        doc.add_paragraph(f"Verification Scan ID: {diff.get('verify_scan_id')}")
        doc.add_paragraph(f"Resolved Findings: {len(diff.get('resolved', []))}")
        doc.add_paragraph(f"Persisted Findings: {len(diff.get('persisted', []))}")
        doc.add_paragraph(f"New Findings: {len(diff.get('new_findings', []))}")
    else:
        doc.add_paragraph("[Populated after post-remediation verification scan]")

    doc.add_heading("3. Audit Log (Remediation Activity)", level=3)
    if logs:
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light List Accent 1"
        h = tbl.rows[0].cells
        for i, label in enumerate(["Timestamp", "Plugin", "Severity", "Command", "Decision", "Status"]):
            h[i].text = label
        for row in logs:
            _table_row(tbl,
                       str(row.get("timestamp", ""))[:19],
                       row.get("plugin_name", ""),
                       row.get("severity_level", ""),
                       row.get("command_dispatched", ""),
                       row.get("operator_decision", ""),
                       row.get("execution_status", ""))
    else:
        doc.add_paragraph("[Populated after remediation run completes]")

    return _save(doc, "Verification_Report.docx")


# ── 8. Final Remediation Report ───────────────────────────────────────────────

def gen_final_report(logs=None, diff=None):
    doc = Document()
    _header(doc, "Final Remediation Report")

    doc.add_heading("1. Project Summary", level=3)
    doc.add_paragraph(
        "AutoRemediate AI is an autonomous vulnerability remediation system developed as "
        "DEPI Capstone Project 7. It integrates with Nessus Professional to perform live "
        "vulnerability scans, applies a multi-factor priority scoring algorithm, presents "
        "human-in-the-loop decision cards for all Critical/High findings, executes approved "
        "remediations over SSH, verifies with two independent checks, and maintains a full "
        "audit trail in Supabase."
    )

    doc.add_heading("2. Metrics", level=3)
    if logs:
        success = sum(1 for r in logs if r.get("execution_status") == "Success")
        failed = sum(1 for r in logs if r.get("execution_status") == "Failed")
        skipped = sum(1 for r in logs if r.get("execution_status") == "Skipped")
        total = len(logs)
        doc.add_paragraph(f"Total decisions logged: {total}")
        doc.add_paragraph(f"Successful remediations: {success}")
        doc.add_paragraph(f"Failed remediations: {failed}")
        doc.add_paragraph(f"Rejected/Skipped: {skipped}")
        if diff:
            doc.add_paragraph(f"Findings resolved: {len(diff.get('resolved', []))}")
    else:
        doc.add_paragraph("[Metrics populated after remediation run]")

    doc.add_heading("3. GitHub Repository", level=3)
    doc.add_paragraph("https://github.com/nourinsayed55-ui/autoremediate-ai")

    doc.add_heading("4. Architecture Overview", level=3)
    for layer, desc in [
        ("Ingestion", "NessusClient (scanner_client.py) calls live Nessus API; Parser normalises findings"),
        ("Orchestration", "Engine (engine.py) scores findings via Recommender, presents Decision Cards, blocks on human input"),
        ("Remediation Transport", "Remediator (remediator.py) executes approved commands over Paramiko SSH with Upstart-safe commands"),
        ("Persistence", "DB (db.py) writes Vulnerabilities and Remediation_Activity_Logs to Supabase"),
        ("Presentation Console", "CLI with colour severity banners, decision cards, and final summary table"),
    ]:
        doc.add_paragraph(f"• {layer}: {desc}")

    doc.add_heading("5. Lessons Learned", level=3)
    for lesson in [
        "Upstart vs systemd: Metasploitable2 requires 'service X stop' not 'systemctl stop X'.",
        "SSH safety must be enforced in code, not just policy — assert_not_ssh_kill() blocks stops and iptables rules for port 22.",
        "Two-factor verification (on-target + external) prevents false positives when the stop command exits 0 but the service restarts.",
        "Supabase's hosted PostgreSQL requires sslmode=require in the connection string.",
        "Human-in-the-loop is non-negotiable: no silent auto-fire path exists even when require_manual_approval=false.",
    ]:
        doc.add_paragraph(f"• {lesson}")

    return _save(doc, "Final_Remediation_Report.docx")


# ── master entry point ────────────────────────────────────────────────────────

def generate_all(cfg=None, scan_id=None, verify_diff_path=None):
    """Generate all docs. Pass findings if scan is available."""
    print("[Docs] Generating documentation …")
    findings = None
    diff = None
    logs = None

    # Load findings if scan results are available
    if scan_id:
        json_path = os.path.join(DOCS_DIR, f"scan_{scan_id}_results.json")
        if os.path.exists(json_path):
            import sys as _sys
            _sys.path.insert(0, os.path.dirname(__file__))
            from parser import load_from_json_file
            from recommender import prioritise_and_recommend
            findings = load_from_json_file(json_path)
            findings = prioritise_and_recommend(findings)
            print(f"[Docs] Loaded {len(findings)} findings from {json_path}")

    # Load verification diff
    if verify_diff_path and os.path.exists(verify_diff_path):
        with open(verify_diff_path) as f:
            diff = json.load(f)

    # Load logs from DB
    try:
        import db as _db
        logs = _db.fetch_all_logs()
    except Exception as exc:
        print(f"[Docs] Could not load DB logs (non-fatal): {exc}")

    gen_tool_config()
    gen_assessment_scope()
    gen_vuln_scan_report(findings, scan_id)
    gen_initial_analysis(findings)
    gen_prioritization_report(findings)
    gen_remediation_plan(findings)
    gen_verification_report(diff, logs)
    gen_final_report(logs, diff)
    print("[Docs] All documents generated.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate AutoRemediate AI documentation")
    parser.add_argument("--scan-id", type=int, default=None)
    parser.add_argument("--verify-diff", default=None, metavar="PATH")
    args = parser.parse_args()
    generate_all(scan_id=args.scan_id, verify_diff_path=args.verify_diff)
